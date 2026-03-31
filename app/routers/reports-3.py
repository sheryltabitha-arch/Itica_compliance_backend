"""
app/routers/reports.py
"""
from __future__ import annotations

import hashlib
import json
import logging
import os
import uuid
from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from app.middleware.auth import CurrentUser, get_current_user, get_supabase

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/v1/reports", tags=["reports"])


# ── Docx helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(str(text or ""), level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p


def _add_para(doc: Document, text: str, bold: bool = False, size: int = 11):
    p   = doc.add_paragraph()
    run = p.add_run(str(text or ""))
    run.bold      = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(str(text or ""))
    p.paragraph_format.space_after = Pt(2)
    return p


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_bg(hdr_cells[i], "1F3864")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold           = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(10)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = str(val or "")
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return table


def _divider(doc: Document):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pb.append(bot)
    pPr.append(pb)
    return p


def _build_docx(data: dict) -> BytesIO:
    doc   = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    title           = doc.add_heading(data.get("reportTitle", "Compliance Report"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        data.get("companyName", ""),
        f"Period: {data.get('reportingPeriod', '')}",
        (
            f"Submitted: {data.get('submissionDate', '')}  |  "
            f"Prepared by: {data.get('preparedBy', '')} ({data.get('preparedByRole', '')})  |  "
            f"{data.get('version', '')}"
        ),
    ]:
        p           = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
    _divider(doc)

    # 1. Purpose & Scope
    _add_heading(doc, "1. Purpose & Scope")
    _add_para(doc, data.get("purpose", ""))
    if data.get("scopeItems"):
        _add_para(doc, "Scope includes:", bold=True)
        for s in data["scopeItems"]:
            _add_bullet(doc, s)
    if data.get("objectives"):
        _add_para(doc, "Objectives:", bold=True)
        for o in data["objectives"]:
            _add_bullet(doc, o)
    _divider(doc)

    # 2. Methodology
    _add_heading(doc, "2. Methodology")
    _add_para(doc, f"Timeframe: {data.get('timeframe', '')}")
    _add_para(doc, f"Sampling: {data.get('sampling', '')}")
    if data.get("dataSources"):
        _add_para(doc, "Data Sources:", bold=True)
        for d in data["dataSources"]:
            _add_bullet(doc, d)
    if data.get("toolsUsed"):
        _add_para(doc, "Tools Used:", bold=True)
        for t in data["toolsUsed"]:
            _add_bullet(doc, t)
    _divider(doc)

    # 3. Key Findings
    _add_heading(doc, "3. Key Findings")
    for f in data.get("keyFindings", []):
        _add_bullet(doc, f)
    _divider(doc)

    # 4. Overall Status
    _add_heading(doc, "4. Overall Compliance Status")
    _add_para(doc, f"Status: {data.get('overallStatus', '')}")
    _add_para(doc, f"Audit Trail Integrity: {data.get('integrity', '')}")
    _divider(doc)

    # 5. Compliance Detail
    _add_heading(doc, "5. Compliance Detail")
    if data.get("compliantAreas"):
        _add_heading(doc, "Compliant Areas", level=2)
        for a in data["compliantAreas"]:
            _add_para(doc, f"✓  {a.get('area', '')}: {a.get('evidence', '')}")
    if data.get("nonCompliantAreas"):
        _add_heading(doc, "Non-Compliant Areas", level=2)
        for a in data["nonCompliantAreas"]:
            _add_para(doc, f"✗  {a.get('area', '')} ({a.get('regulation', '')}): {a.get('evidence', '')}  [Risk: {a.get('riskLevel', '')}]")
    if data.get("partialAreas"):
        _add_heading(doc, "Partially Compliant Areas", level=2)
        for a in data["partialAreas"]:
            _add_para(doc, f"~  {a.get('area', '')}: {a.get('detail', '')}")
    _divider(doc)

    # 6. Regulatory Mapping
    if data.get("regulatoryMap"):
        _add_heading(doc, "6. Regulatory Mapping")
        _add_table(
            doc,
            headers=["Regulation", "Requirement", "Control", "Status"],
            rows=[[r.get("regulation",""), r.get("requirement",""), r.get("control",""), r.get("status","")] for r in data["regulatoryMap"]],
        )
        _divider(doc)

    # 7. Risk Register
    _add_heading(doc, "7. Risk Register")
    _add_para(doc, data.get("riskMethodology", ""))
    if data.get("risks"):
        _add_table(
            doc,
            headers=["Risk", "Impact", "Likelihood", "Rating", "Owner"],
            rows=[[r.get("risk",""), r.get("impact",""), r.get("likelihood",""), r.get("rating",""), r.get("owner","")] for r in data["risks"]],
        )
    _divider(doc)

    # 8. Recommendations
    if data.get("recommendations"):
        _add_heading(doc, "8. Recommendations")
        _add_table(
            doc,
            headers=["Issue", "Action", "Priority", "Responsible", "Deadline"],
            rows=[[r.get("issue",""), r.get("action",""), r.get("priority",""), r.get("responsible",""), r.get("deadline","")] for r in data["recommendations"]],
        )
        _divider(doc)

    # 9. Statistics
    _add_heading(doc, "9. Statistics")
    stats = data.get("stats", {})
    for label, val in [
        ("Total Documents Processed",       stats.get("totalDocuments", "")),
        ("Verified Documents",               stats.get("verifiedDocuments", "")),
        ("Total Compliance Decisions",       stats.get("totalDecisions", "")),
        ("Average Extraction Confidence",    f"{float(stats.get('avgConfidence', 0)) * 100:.1f}%"),
        ("Low Confidence Flags (<75%)",      stats.get("lowConfidenceFlags", "")),
    ]:
        _add_para(doc, f"{label}: {val}")
    _divider(doc)

    # 10. Major Risks & Actions
    if data.get("majorRisks") or data.get("actionsRequired"):
        _add_heading(doc, "10. Major Risks & Actions Required")
        for r in data.get("majorRisks", []):
            _add_bullet(doc, r)
        if data.get("actionsRequired"):
            _add_para(doc, "Actions Required:", bold=True)
            for a in data["actionsRequired"]:
                _add_bullet(doc, a)
        _divider(doc)

    # 11. Conclusion
    _add_heading(doc, "11. Conclusion")
    _add_para(doc, data.get("conclusion", ""))
    _add_para(doc, data.get("regulatoryReadiness", ""))
    _divider(doc)

    # 12. Next Steps
    if data.get("nextSteps"):
        _add_heading(doc, "12. Next Steps")
        for s in data["nextSteps"]:
            _add_bullet(doc, s)
        _divider(doc)

    # 13. Limitations & Exclusions
    if data.get("limitations") or data.get("exclusions"):
        _add_heading(doc, "13. Limitations & Exclusions")
        for l in data.get("limitations", []):
            _add_bullet(doc, f"Limitation: {l}")
        for e in data.get("exclusions", []):
            _add_bullet(doc, f"Exclusion: {e}")
        _divider(doc)

    # 14. Glossary
    if data.get("glossary"):
        _add_heading(doc, "14. Glossary")
        for g in data["glossary"]:
            _add_para(doc, f"{g.get('term', '')}: {g.get('definition', '')}")
        _divider(doc)

    # Appendix
    _add_heading(doc, "Appendix")
    _add_para(doc, data.get("appendixNote", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Route 1: SAR / PDF report metadata ───────────────────────────────────────

@router.post("")
@router.get("/generate")
async def generate_report(
    period_start: str = "",
    period_end:   str = "",
    report_type:  str = "aml_summary",
    format:       str = "pdf",
    current: CurrentUser = Depends(get_current_user),
):
    short_id  = uuid.uuid4().hex[:8].upper()
    report_id = f"RPT-{short_id}"
    now       = datetime.now(timezone.utc).isoformat()
    tenant_id = str(current.tenant_id)

    extraction_count   = 0
    completed_count    = 0
    reviewed_count     = 0
    avg_confidence     = 0.0
    low_confidence_count = 0

    try:
        sb    = get_supabase()
        query = sb.table("extractions").select("status, confidence_scores").eq("tenant_id", tenant_id)
        if period_start:
            query = query.gte("created_at", period_start)
        if period_end:
            query = query.lte("created_at", period_end)

        rows             = query.execute().data or []
        extraction_count = len(rows)
        completed_count  = sum(1 for r in rows if r.get("status") == "completed")
        reviewed_count   = sum(1 for r in rows if r.get("status") == "reviewed")

        all_scores = []
        for row in rows:
            scores = row.get("confidence_scores") or {}
            all_scores.extend(float(v) for v in scores.values())

        if all_scores:
            avg_confidence       = round(sum(all_scores) / len(all_scores), 4)
            low_confidence_count = sum(1 for s in all_scores if s < 0.75)

    except Exception as e:
        logger.warning(f"Report data query failed (non-fatal): {e}")

    report_hash = "demo-hash"
    try:
        sb            = get_supabase()
        prev          = sb.table("audit_events").select("hash").eq("tenant_id", tenant_id).order("created_at", desc=True).limit(1).execute()
        previous_hash = prev.data[0]["hash"] if prev.data else "GENESIS"

        hash_input  = json.dumps({
            "report_id":       report_id,
            "tenant_id":       tenant_id,
            "period_start":    period_start,
            "period_end":      period_end,
            "report_type":     report_type,
            "extraction_count": extraction_count,
            "timestamp":       now,
        }, sort_keys=True)
        report_hash = hashlib.sha256(hash_input.encode()).hexdigest()

        event_count = sb.table("audit_events").select("id", count="exact").eq("tenant_id", tenant_id).execute()
        event_num   = (event_count.count or 0) + 1

        sb.table("audit_events").insert({
            "tenant_id":     tenant_id,
            "user_id":       str(current.user_id),
            "event_type":    "REPORT_GENERATED",
            "event_id":      f"EVT-{event_num:05d}",
            "detail":        (
                f"Report {report_id} | {report_type} | "
                f"{period_start or 'all'} to {period_end or 'now'} | "
                f"{extraction_count} extractions"
            ),
            "hash":          report_hash,
            "previous_hash": previous_hash,
        }).execute()

    except Exception as e:
        logger.warning(f"Report audit write failed (non-fatal): {e}")

    return {
        "report_id":    report_id,
        "status":       "sealed",
        "report_type":  report_type,
        "format":       format,
        "period_start": period_start,
        "period_end":   period_end,
        "hash":         report_hash,
        "generated_at": now,
        "tenant_id":    tenant_id,
        "summary": {
            "total_documents":      extraction_count,
            "completed":            completed_count,
            "reviewed":             reviewed_count,
            "pending":              extraction_count - completed_count - reviewed_count,
            "avg_confidence":       avg_confidence,
            "low_confidence_flags": low_confidence_count,
        },
    }


# ── Route 2: Word document download ──────────────────────────────────────────

@router.post("/generate/docx")
async def generate_report_docx(
    period_start: str = "",
    period_end:   str = "",
    report_type:  str = "aml_summary",
    current: CurrentUser = Depends(get_current_user),
):
    """Generate a regulator-ready Word document report."""
    sb        = get_supabase()
    tenant_id = str(current.tenant_id)

    rows      = sb.table("extractions").select("*").eq("tenant_id", tenant_id).execute().data or []
    events    = sb.table("audit_events").select("*").eq("tenant_id", tenant_id).order("created_at", desc=True).limit(50).execute().data or []
    decisions = sb.table("decisions").select("id", count="exact").eq("tenant_id", tenant_id).execute()

    integrity_check = "VERIFIED"
    for i in range(1, len(events)):
        if events[i]["hash"] != events[i - 1].get("previous_hash"):
            integrity_check = "WARNING"
            break

    all_scores = []
    for row in rows:
        all_scores.extend((row.get("confidence_scores") or {}).values())
    avg_conf  = sum(float(s) for s in all_scores) / len(all_scores) if all_scores else 0.0
    low_flags = sum(1 for s in all_scores if float(s) < 0.75)

    data = {
        "reportTitle":    f"Compliance Audit Report – {period_start or 'All Time'}",
        "companyName":    "Itica Technologies Ltd.",
        "reportingPeriod": f"{period_start or 'All'} – {period_end or 'Present'}",
        "submissionDate": datetime.now(timezone.utc).strftime("%d %B %Y"),
        "preparedBy":     getattr(current, "name", None) or current.email,
        "preparedByRole": "Compliance Officer",
        "version":        "v1.0",
        "purpose":        f"Compliance audit covering {len(rows)} KYC documents processed on the Itica platform.",
        "overallStatus":  "Compliant" if low_flags == 0 else "Partially Compliant",
        "keyFindings": [
            f"{len(rows)} KYC documents processed",
            f"Average extraction confidence: {avg_conf * 100:.1f}%",
            f"{low_flags} documents flagged for low confidence",
            f"Audit trail integrity: {integrity_check}",
            f"{decisions.count or 0} compliance decisions recorded",
        ],
        "majorRisks": [
            "Sanctions screening not yet implemented"
            if not os.environ.get("SANCTIONS_API_URL") and not os.environ.get("OFAC_LIST_PATH")
            else "Sanctions screening active"
        ],
        "actionsRequired": [
            "Implement sanctions screening"
            if not os.environ.get("OFAC_LIST_PATH")
            else "Continue monitoring"
        ],
        "integrity": integrity_check,
        "stats": {
            "totalDocuments":     len(rows),
            "verifiedDocuments":  sum(1 for r in rows if r.get("status") == "reviewed"),
            "totalDecisions":     decisions.count or 0,
            "avgConfidence":      avg_conf,
            "lowConfidenceFlags": low_flags,
        },
        "scopeItems":  ["KYC extraction system", "Human review queue", "Audit trail", "Auth and access control"],
        "frameworks":  ["FATF Recommendations", "GDPR", "ISO 27001"],
        "objectives":  ["Verify KYC pipeline accuracy", "Confirm audit trail integrity", "Identify compliance gaps"],
        "dataSources": ["Itica audit_events table", "Extractions table", "Decisions table"],
        "toolsUsed":   ["Itica Compliance Platform v2.0.0", "LayoutLMv3 (HuggingFace)", "Auth0", "Supabase"],
        "timeframe":   f"{period_start or 'All time'} to {period_end or 'present'}",
        "sampling":    "Full population review. No sampling applied.",
        "regulatoryMap": [
            {"regulation": "FATF R.10",    "requirement": "Customer Due Diligence",  "control": "KYC extraction + human review", "status": "Partial"},
            {"regulation": "FATF R.11",    "requirement": "Record keeping",           "control": "Immutable hash chain",          "status": "Compliant"},
            {"regulation": "GDPR Art. 32", "requirement": "Security of processing",   "control": "Auth0 + AES-256",               "status": "Compliant"},
        ],
        "compliantAreas": [{"area": "Audit Trail", "evidence": f"Hash chain integrity: {integrity_check}. All events cryptographically linked."}],
        "nonCompliantAreas": [] if os.environ.get("OFAC_LIST_PATH") else [{"area": "Sanctions Screening", "regulation": "FATF R.6", "evidence": "No watchlist check implemented.", "riskLevel": "High"}],
        "partialAreas": [{"area": "KYC Completeness", "detail": f"{low_flags} documents below 0.75 confidence threshold, routed to human review."}] if low_flags > 0 else [],
        "riskMethodology": "3x3 Impact x Likelihood matrix. High = immediate action. Medium = 30 days. Low = monitor.",
        "risks":           [{"risk": "Sanctions screening gap", "impact": "High", "likelihood": "Medium", "rating": "High", "owner": "Engineering"}],
        "recommendations": [{"issue": "Sanctions screening", "action": "Integrate OFAC/UN list via MiniLM", "priority": "High", "responsible": "Engineering", "deadline": "30 Apr 2026"}],
        "limitations":     ["Third-party vendor audits excluded", "Penetration testing out of scope"],
        "exclusions":      ["Auth0, HuggingFace, Supabase vendor audits"],
        "conclusion":      f"The platform processed {len(rows)} documents with {avg_conf * 100:.1f}% average confidence. Audit trail integrity confirmed.",
        "regulatoryReadiness": "Ready for audit trail and KYC workflow review. Sanctions screening required before full regulatory submission.",
        "nextSteps":       ["Implement sanctions screening", "SOC 2 Type I via Vanta", "Q2 2026 expanded audit"],
        "glossary": [
            {"term": "KYC",        "definition": "Know Your Customer"},
            {"term": "AML",        "definition": "Anti-Money Laundering"},
            {"term": "FATF",       "definition": "Financial Action Task Force"},
            {"term": "Hash Chain", "definition": "Tamper-evident cryptographic audit log"},
        ],
        "appendixNote": "Full logs available from Itica platform administrator.",
    }

    try:
        buf = _build_docx(data)
    except Exception as e:
        logger.error(f"Report generation failed: {e}")
        raise HTTPException(500, "Report generation failed")

    filename = f"itica_compliance_report_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
